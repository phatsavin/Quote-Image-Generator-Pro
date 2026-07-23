from __future__ import annotations

import csv
import tempfile
import unittest
from dataclasses import replace
from pathlib import Path

from docx import Document
from PIL import Image, ImageChops

from quote_engine import (
    QuoteGeneratorError,
    RenderSettings,
    create_generation_items,
    generate_batch,
    list_images,
    load_quotes,
    render_quote_image,
)


class QuoteEngineTests(unittest.TestCase):
    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory()
        self.root = Path(self.temp_dir.name)
        self.images_dir = self.root / "images"
        self.images_dir.mkdir()
        for index, color in enumerate(
            [(20, 55, 90), (120, 68, 48), (38, 96, 75)], start=1
        ):
            image = Image.new("RGB", (900 + index * 30, 700 + index * 50), color)
            image.save(self.images_dir / f"background_{index}.jpg")

    def tearDown(self):
        self.temp_dir.cleanup()

    def test_load_txt_and_clean_numbering(self):
        path = self.root / "quotes.txt"
        path.write_text(
            '1. "First original quote."\n2) “Second original quote.”\n\n'
            "First original quote.\n",
            encoding="utf-8",
        )
        self.assertEqual(
            load_quotes(path),
            ["First original quote.", "Second original quote."],
        )

    def test_load_csv(self):
        path = self.root / "quotes.csv"
        with path.open("w", newline="", encoding="utf-8") as handle:
            writer = csv.writer(handle)
            writer.writerow(["Quote"])
            writer.writerow(["A useful quote for testing."])
            writer.writerow(["Another useful quote for testing."])
        self.assertEqual(len(load_quotes(path)), 2)

    def test_load_docx_quote_item_style(self):
        path = self.root / "quotes.docx"
        document = Document()
        styles = document.styles
        styles.add_style("Quote Item", 1)
        document.add_paragraph("Document title", style="Title")
        document.add_paragraph('"First DOCX quote."', style="Quote Item")
        document.add_paragraph('"Second DOCX quote."', style="Quote Item")
        document.save(path)
        self.assertEqual(
            load_quotes(path),
            ["First DOCX quote.", "Second DOCX quote."],
        )

    def test_create_items_without_repeats(self):
        images = list_images(self.images_dir)
        quotes = ["One quote", "Two quote", "Three quote"]
        items = create_generation_items(
            images,
            quotes,
            3,
            random_images=True,
            random_quotes=True,
            seed=42,
        )
        self.assertEqual(len({item.image_path for item in items}), 3)
        self.assertEqual(len({item.quote for item in items}), 3)

    def test_repeated_item_guard(self):
        with self.assertRaises(QuoteGeneratorError):
            create_generation_items(
                list_images(self.images_dir),
                ["Only quote"],
                2,
                avoid_repeat_quotes=True,
            )

    def test_render_dimensions(self):
        settings = RenderSettings(width=600, height=750, font_max_size=54)
        output = render_quote_image(
            self.images_dir / "background_1.jpg",
            "A calm and readable quote for a Facebook image.",
            settings,
        )
        self.assertEqual(output.size, (600, 750))
        self.assertEqual(output.mode, "RGB")

    def test_background_styles_render_distinct_results(self):
        settings = RenderSettings(
            width=360,
            height=450,
            font_min_size=22,
            font_max_size=36,
            darken_percent=0,
            box_enabled=False,
        )
        styles = [
            "classic",
            "tiktok-blur-frame",
            "tiktok-glass-gradient",
            "cinematic-vignette",
            "luxury-noir",
            "luxury-noir-color",
        ]
        fingerprints = []
        for style in styles:
            output = render_quote_image(
                self.images_dir / "background_2.jpg",
                "Professional quote style preview.",
                replace(settings, background_style=style),
            )
            self.assertEqual(output.size, (360, 450))
            fingerprints.append(output.resize((16, 20)).tobytes())
        self.assertEqual(len(set(fingerprints)), len(styles))

    def test_luxury_noir_color_preserves_source_color(self):
        settings = RenderSettings(
            width=360,
            height=450,
            font_min_size=22,
            font_max_size=36,
            background_style="luxury-noir-color",
            darken_percent=0,
            box_enabled=False,
            shadow_enabled=False,
        )
        output = render_quote_image(
            self.images_dir / "background_2.jpg",
            "Original color frame test.",
            settings,
        )
        source_color = Image.open(
            self.images_dir / "background_2.jpg"
        ).convert("RGB").getpixel((10, 10))
        self.assertEqual(output.getpixel((180, 55)), source_color)

    def test_luxury_noir_color_uses_thick_gold_frame(self):
        settings = RenderSettings(
            width=600,
            height=750,
            font_min_size=22,
            font_max_size=36,
            background_style="luxury-noir-color",
            darken_percent=0,
            box_enabled=False,
            shadow_enabled=False,
        )
        framed = render_quote_image(
            self.images_dir / "background_2.jpg",
            "Thick frame test.",
            settings,
        )
        classic = render_quote_image(
            self.images_dir / "background_2.jpg",
            "Thick frame test.",
            replace(settings, background_style="classic"),
        )
        difference = ImageChops.difference(classic, framed)
        inset = max(10, round(min(settings.width, settings.height) * 0.016))
        changed_rows = [
            y
            for y in range(inset, inset + 8)
            if difference.getpixel((settings.width // 2, y)) != (0, 0, 0)
        ]
        self.assertGreaterEqual(len(changed_rows), 3)

    def test_tiktok_glass_renders_with_quote_box(self):
        output = render_quote_image(
            self.images_dir / "background_1.jpg",
            "Glass quote box rendering test.",
            RenderSettings(
                width=360,
                height=450,
                font_min_size=22,
                font_max_size=36,
                background_style="tiktok-glass-gradient",
                box_enabled=True,
                box_opacity=100,
            ),
        )
        self.assertEqual(output.size, (360, 450))

    def test_quote_frame_has_outline_and_no_fill(self):
        settings = RenderSettings(
            width=500,
            height=625,
            font_min_size=30,
            font_max_size=48,
            quote_position="center",
            add_quote_marks=False,
            darken_percent=0,
            box_enabled=True,
            box_style="quote-frame",
            box_color="#101010",
            box_opacity=220,
            shadow_enabled=False,
        )
        framed = render_quote_image(
            self.images_dir / "background_1.jpg",
            "Frame only",
            settings,
        )
        plain = render_quote_image(
            self.images_dir / "background_1.jpg",
            "Frame only",
            replace(settings, box_enabled=False),
        )
        difference = ImageChops.difference(plain, framed).convert("L")
        bbox = difference.getbbox()
        self.assertIsNotNone(bbox)
        assert bbox is not None
        cropped = difference.crop(bbox)
        histogram = cropped.histogram()
        changed_pixels = sum(histogram[1:])
        total_pixels = cropped.width * cropped.height
        self.assertGreater(changed_pixels / total_pixels, 0.02)
        self.assertLess(changed_pixels / total_pixels, 0.35)

    def test_corner_quote_frame_has_corners_and_no_fill(self):
        settings = RenderSettings(
            width=500,
            height=625,
            font_min_size=30,
            font_max_size=48,
            quote_position="center",
            add_quote_marks=False,
            darken_percent=0,
            box_enabled=True,
            box_style="corner-frame",
            box_color="#101010",
            box_opacity=220,
            shadow_enabled=False,
        )
        framed = render_quote_image(
            self.images_dir / "background_1.jpg",
            "Corner frame",
            settings,
        )
        plain = render_quote_image(
            self.images_dir / "background_1.jpg",
            "Corner frame",
            replace(settings, box_enabled=False),
        )
        difference = ImageChops.difference(plain, framed).convert("L")
        bbox = difference.getbbox()
        self.assertIsNotNone(bbox)
        assert bbox is not None
        cropped = difference.crop(bbox)
        histogram = cropped.histogram()
        changed_pixels = sum(histogram[1:])
        total_pixels = cropped.width * cropped.height
        self.assertGreater(changed_pixels / total_pixels, 0.02)
        self.assertLess(changed_pixels / total_pixels, 0.30)

    def test_double_line_frame_has_two_outlines_and_no_fill(self):
        settings = RenderSettings(
            width=500,
            height=625,
            font_min_size=30,
            font_max_size=48,
            quote_position="center",
            add_quote_marks=True,
            darken_percent=0,
            box_enabled=True,
            box_style="double-frame",
            box_color="#101010",
            box_opacity=220,
            shadow_enabled=False,
        )
        framed = render_quote_image(
            self.images_dir / "background_1.jpg",
            "Double line frame",
            settings,
        )
        plain = render_quote_image(
            self.images_dir / "background_1.jpg",
            "Double line frame",
            replace(settings, box_enabled=False),
        )
        difference = ImageChops.difference(plain, framed).convert("L")
        bbox = difference.getbbox()
        self.assertIsNotNone(bbox)
        assert bbox is not None
        cropped = difference.crop(bbox)
        histogram = cropped.histogram()
        changed_pixels = sum(histogram[1:])
        total_pixels = cropped.width * cropped.height
        self.assertGreater(changed_pixels / total_pixels, 0.02)
        self.assertLess(changed_pixels / total_pixels, 0.30)

    def test_watermark_top_and_bottom_positions(self):
        base_settings = RenderSettings(
            width=400,
            height=500,
            font_min_size=22,
            font_max_size=34,
            darken_percent=0,
            box_enabled=False,
            shadow_enabled=False,
            watermark="",
        )
        base = render_quote_image(
            self.images_dir / "background_3.jpg",
            "Watermark position test.",
            base_settings,
        )
        positions = {
            "top-left": ("top", "left"),
            "top-center": ("top", "center"),
            "top-right": ("top", "right"),
            "bottom-left": ("bottom", "left"),
            "bottom-center": ("bottom", "center"),
            "bottom-right": ("bottom", "right"),
        }
        for position, (vertical, horizontal) in positions.items():
            output = render_quote_image(
                self.images_dir / "background_3.jpg",
                "Watermark position test.",
                replace(
                    base_settings,
                    watermark="Life Quote",
                    watermark_position=position,
                ),
            )
            difference = ImageChops.difference(base, output)
            bbox = difference.getbbox()
            self.assertIsNotNone(bbox)
            assert bbox is not None
            center_x = (bbox[0] + bbox[2]) / 2
            center_y = (bbox[1] + bbox[3]) / 2
            if vertical == "top":
                self.assertLess(center_y, base_settings.height / 3)
            else:
                self.assertGreater(center_y, base_settings.height * 2 / 3)
            if horizontal == "left":
                self.assertLess(center_x, base_settings.width / 3)
            elif horizontal == "center":
                self.assertGreater(center_x, base_settings.width / 3)
                self.assertLess(center_x, base_settings.width * 2 / 3)
            else:
                self.assertGreater(center_x, base_settings.width * 2 / 3)

    def test_invalid_background_and_watermark_settings(self):
        with self.assertRaises(QuoteGeneratorError):
            RenderSettings(background_style="unknown").validate()
        with self.assertRaises(QuoteGeneratorError):
            RenderSettings(watermark_position="middle").validate()
        with self.assertRaises(QuoteGeneratorError):
            RenderSettings(box_style="filled-outline").validate()

    def test_generate_batch_and_log(self):
        images = list_images(self.images_dir)
        progress_paths = []
        items = create_generation_items(
            images,
            ["Quote one", "Quote two"],
            2,
            random_images=False,
            random_quotes=False,
        )
        folder, results = generate_batch(
            items,
            self.root / "output",
            RenderSettings(width=500, height=625, font_max_size=48),
            progress_callback=lambda current, total, result: progress_paths.append(
                (current, total, result.output_path)
            ),
            create_timestamp_folder=False,
        )
        self.assertEqual(len(results), 2)
        self.assertEqual(len(progress_paths), 2)
        self.assertEqual(
            [current for current, _, _ in progress_paths],
            [1, 2],
        )
        self.assertTrue(all(path.is_file() for _, _, path in progress_paths))
        self.assertTrue((folder / "generation_log.csv").is_file())
        for result in results:
            self.assertTrue(result.output_path.is_file())
            with Image.open(result.output_path) as generated:
                self.assertEqual(generated.size, (500, 625))


if __name__ == "__main__":
    unittest.main()
