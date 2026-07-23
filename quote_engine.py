from __future__ import annotations

import csv
import random
import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Callable, Iterable, Sequence

from PIL import Image, ImageDraw, ImageEnhance, ImageFilter, ImageFont, ImageOps


SUPPORTED_IMAGE_EXTENSIONS = {
    ".jpg",
    ".jpeg",
    ".png",
    ".webp",
    ".bmp",
    ".tif",
    ".tiff",
}
SUPPORTED_QUOTE_EXTENSIONS = {".txt", ".csv", ".docx"}
BACKGROUND_STYLES = {
    "classic",
    "tiktok-blur-frame",
    "tiktok-glass-gradient",
    "cinematic-vignette",
    "luxury-noir",
    "luxury-noir-color",
}
WATERMARK_POSITIONS = {
    "top-left",
    "top-center",
    "top-right",
    "bottom-left",
    "bottom-center",
    "bottom-right",
}
QUOTE_BOX_STYLES = {
    "translucent",
    "quote-frame",
    "corner-frame",
    "double-frame",
}


class QuoteGeneratorError(RuntimeError):
    """User-facing error raised by the quote image engine."""


@dataclass(frozen=True)
class RenderSettings:
    width: int = 1440
    height: int = 1800
    output_format: str = "JPEG"
    jpeg_quality: int = 95
    font_path: str = ""
    font_min_size: int = 42
    font_max_size: int = 104
    text_color: str = "#FFFFFF"
    quote_position: str = "left"
    vertical_position: int = 50
    safe_margin_percent: int = 8
    text_area_percent: int = 48
    line_spacing_ratio: float = 0.28
    add_quote_marks: bool = True
    background_style: str = "classic"
    darken_percent: int = 22
    blur_radius: float = 0.0
    box_enabled: bool = True
    box_style: str = "translucent"
    box_color: str = "#06111F"
    box_opacity: int = 105
    box_padding_percent: float = 2.4
    box_radius_percent: float = 1.6
    shadow_enabled: bool = True
    shadow_color: str = "#000000"
    shadow_opacity: int = 170
    shadow_offset: int = 5
    stroke_width: int = 1
    watermark: str = ""
    watermark_color: str = "#FFFFFF"
    watermark_opacity: int = 165
    watermark_position: str = "bottom-right"
    focal_point: str = "center"

    def validate(self) -> None:
        if self.width < 320 or self.height < 320:
            raise QuoteGeneratorError("Output size must be at least 320 x 320.")
        if self.width > 8000 or self.height > 8000:
            raise QuoteGeneratorError("Output size cannot exceed 8000 x 8000.")
        if self.font_min_size < 10 or self.font_max_size < self.font_min_size:
            raise QuoteGeneratorError("Invalid minimum/maximum font size.")
        if self.quote_position not in {"left", "center", "right"}:
            raise QuoteGeneratorError("Quote position must be left, center, or right.")
        if self.focal_point not in {"left", "center", "right"}:
            raise QuoteGeneratorError("Focal point must be left, center, or right.")
        if self.background_style not in BACKGROUND_STYLES:
            raise QuoteGeneratorError("Invalid background style.")
        if self.watermark_position not in WATERMARK_POSITIONS:
            raise QuoteGeneratorError("Invalid watermark position.")
        if self.box_style not in QUOTE_BOX_STYLES:
            raise QuoteGeneratorError("Invalid quote box style.")
        if not 0 <= self.vertical_position <= 100:
            raise QuoteGeneratorError("Vertical position must be between 0 and 100.")
        if not 2 <= self.safe_margin_percent <= 25:
            raise QuoteGeneratorError("Safe margin must be between 2% and 25%.")
        if not 25 <= self.text_area_percent <= 90:
            raise QuoteGeneratorError("Text area must be between 25% and 90%.")
        if not 0 <= self.darken_percent <= 85:
            raise QuoteGeneratorError("Darken value must be between 0% and 85%.")
        if not 0 <= self.box_opacity <= 255:
            raise QuoteGeneratorError("Quote box opacity must be between 0 and 255.")
        if self.output_format.upper() not in {"JPEG", "PNG"}:
            raise QuoteGeneratorError("Output format must be JPEG or PNG.")


@dataclass(frozen=True)
class GenerationItem:
    image_path: Path
    quote: str
    index: int


@dataclass(frozen=True)
class GenerationResult:
    output_path: Path
    image_path: Path
    quote: str
    index: int


def _read_text_file(path: Path) -> str:
    errors: list[str] = []
    for encoding in ("utf-8-sig", "utf-16", "cp1252"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeError as exc:
            errors.append(f"{encoding}: {exc}")
    raise QuoteGeneratorError(
        f"Unable to decode quote file: {path.name}. Tried UTF-8, UTF-16, and CP1252."
    )


def _clean_quote(value: str) -> str:
    value = value.replace("\ufeff", "").strip()
    value = re.sub(r"^\s*\d{1,5}\s*[\.\):\-]\s*", "", value)
    value = re.sub(r"\s+", " ", value).strip()
    pairs = [
        ('"', '"'),
        ("\u201c", "\u201d"),
        ("\u2018", "\u2019"),
        ("'", "'"),
    ]
    for left, right in pairs:
        if len(value) >= 2 and value.startswith(left) and value.endswith(right):
            value = value[len(left) : len(value) - len(right)].strip()
            break
    return value


def _deduplicate(values: Iterable[str]) -> list[str]:
    result: list[str] = []
    seen: set[str] = set()
    for value in values:
        cleaned = _clean_quote(value)
        key = cleaned.casefold()
        if cleaned and key not in seen:
            result.append(cleaned)
            seen.add(key)
    return result


def load_quotes(path_value: str | Path) -> list[str]:
    path = Path(path_value)
    if not path.is_file():
        raise QuoteGeneratorError(f"Quote file does not exist: {path}")
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_QUOTE_EXTENSIONS:
        raise QuoteGeneratorError("Quote file must be TXT, CSV, or DOCX.")

    if suffix == ".txt":
        raw = _read_text_file(path)
        quotes = _deduplicate(line for line in raw.splitlines() if line.strip())
    elif suffix == ".csv":
        raw = _read_text_file(path)
        rows = list(csv.reader(raw.splitlines()))
        values: list[str] = []
        for row in rows:
            if not row:
                continue
            candidate = next((cell for cell in row if cell.strip()), "")
            if candidate.strip().casefold() in {"quote", "quotes", "text", "caption"}:
                continue
            values.append(candidate)
        quotes = _deduplicate(values)
    else:
        try:
            from docx import Document
        except ImportError as exc:
            raise QuoteGeneratorError(
                "DOCX support requires python-docx. Run: pip install python-docx"
            ) from exc

        document = Document(path)
        styled = [
            paragraph.text
            for paragraph in document.paragraphs
            if paragraph.text.strip()
            and paragraph.style
            and paragraph.style.name.casefold() == "quote item"
        ]
        if styled:
            quotes = _deduplicate(styled)
        else:
            excluded = {
                "ready for image posts, captions, stories, and short-form video.",
                "50 original english quotes",
                "facebook content collection",
            }
            candidates: list[str] = []
            for paragraph in document.paragraphs:
                text = paragraph.text.strip()
                if not text or text.casefold() in excluded:
                    continue
                if paragraph.style and paragraph.style.name.casefold() in {
                    "title",
                    "subtitle",
                    "heading 1",
                    "heading 2",
                    "heading 3",
                }:
                    continue
                if len(text.split()) < 4:
                    continue
                candidates.append(text)
            quotes = _deduplicate(candidates)

    if not quotes:
        raise QuoteGeneratorError(f"No usable quotes were found in {path.name}.")
    return quotes


def list_images(folder_value: str | Path) -> list[Path]:
    folder = Path(folder_value)
    if not folder.is_dir():
        raise QuoteGeneratorError(f"Image folder does not exist: {folder}")
    images = sorted(
        (
            path
            for path in folder.iterdir()
            if path.is_file() and path.suffix.lower() in SUPPORTED_IMAGE_EXTENSIONS
        ),
        key=lambda item: item.name.casefold(),
    )
    if not images:
        raise QuoteGeneratorError(
            "No JPG, PNG, WEBP, BMP, or TIFF images were found in the selected folder."
        )
    return images


def create_generation_items(
    images: Sequence[Path],
    quotes: Sequence[str],
    count: int,
    *,
    random_images: bool = True,
    random_quotes: bool = False,
    avoid_repeat_images: bool = True,
    avoid_repeat_quotes: bool = True,
    seed: int | None = None,
) -> list[GenerationItem]:
    if count <= 0:
        raise QuoteGeneratorError("Generation count must be greater than zero.")
    if not images:
        raise QuoteGeneratorError("No source images are available.")
    if not quotes:
        raise QuoteGeneratorError("No quotes are available.")
    if avoid_repeat_images and count > len(images):
        raise QuoteGeneratorError(
            f"Requested {count} outputs but the folder contains only {len(images)} images. "
            "Reduce the count or allow repeated images."
        )
    if avoid_repeat_quotes and count > len(quotes):
        raise QuoteGeneratorError(
            f"Requested {count} outputs but the file contains only {len(quotes)} quotes. "
            "Reduce the count or allow repeated quotes."
        )

    rng = random.Random(seed)

    if avoid_repeat_images:
        chosen_images = list(images)
        if random_images:
            rng.shuffle(chosen_images)
        chosen_images = chosen_images[:count]
    else:
        if random_images:
            chosen_images = [rng.choice(images) for _ in range(count)]
        else:
            chosen_images = [images[index % len(images)] for index in range(count)]

    if avoid_repeat_quotes:
        chosen_quotes = list(quotes)
        if random_quotes:
            rng.shuffle(chosen_quotes)
        chosen_quotes = chosen_quotes[:count]
    else:
        if random_quotes:
            chosen_quotes = [rng.choice(quotes) for _ in range(count)]
        else:
            chosen_quotes = [quotes[index % len(quotes)] for index in range(count)]

    return [
        GenerationItem(image_path=image, quote=quote, index=index + 1)
        for index, (image, quote) in enumerate(zip(chosen_images, chosen_quotes))
    ]


def _hex_to_rgba(value: str, alpha: int = 255) -> tuple[int, int, int, int]:
    cleaned = value.strip().lstrip("#")
    if len(cleaned) == 3:
        cleaned = "".join(char * 2 for char in cleaned)
    if not re.fullmatch(r"[0-9a-fA-F]{6}", cleaned):
        raise QuoteGeneratorError(f"Invalid color value: {value}")
    return (
        int(cleaned[0:2], 16),
        int(cleaned[2:4], 16),
        int(cleaned[4:6], 16),
        max(0, min(255, alpha)),
    )


def find_default_font(bold: bool = True) -> str:
    file_names = (
        ["georgiab.ttf", "calibrib.ttf", "arialbd.ttf", "timesbd.ttf"]
        if bold
        else ["georgia.ttf", "calibri.ttf", "arial.ttf", "times.ttf"]
    )
    candidates: list[Path] = []
    for name in file_names:
        candidates.extend(
            [
                Path("C:/Windows/Fonts") / name,
                Path("/usr/share/fonts/truetype/dejavu") / name.replace(
                    "georgiab", "DejaVuSerif-Bold"
                ).replace("georgia", "DejaVuSerif"),
                Path("/usr/share/fonts/truetype/liberation2")
                / (
                    "LiberationSerif-Bold.ttf"
                    if bold
                    else "LiberationSerif-Regular.ttf"
                ),
            ]
        )
    candidates.extend(
        [
            Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
            Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
        ]
    )
    for candidate in candidates:
        if candidate.is_file():
            return str(candidate)
    return ""


def _load_font(path_value: str, size: int, *, bold_fallback: bool = True):
    requested = Path(path_value) if path_value else None
    if requested and requested.is_file():
        try:
            return ImageFont.truetype(str(requested), size=size)
        except OSError as exc:
            raise QuoteGeneratorError(f"Unable to load font: {requested}") from exc

    fallback = find_default_font(bold=bold_fallback)
    if fallback:
        return ImageFont.truetype(fallback, size=size)
    return ImageFont.load_default()


def _measure(draw: ImageDraw.ImageDraw, text: str, font) -> float:
    left, _, right, _ = draw.textbbox((0, 0), text, font=font)
    return right - left


def _split_long_token(
    draw: ImageDraw.ImageDraw, token: str, font, max_width: int
) -> list[str]:
    parts: list[str] = []
    current = ""
    for char in token:
        test = current + char
        if current and _measure(draw, test, font) > max_width:
            parts.append(current)
            current = char
        else:
            current = test
    if current:
        parts.append(current)
    return parts


def wrap_text(
    draw: ImageDraw.ImageDraw, text: str, font, max_width: int
) -> str:
    paragraphs = text.splitlines() or [text]
    lines: list[str] = []
    for paragraph in paragraphs:
        words = paragraph.split()
        if not words:
            lines.append("")
            continue
        current = ""
        for word in words:
            if _measure(draw, word, font) > max_width:
                chunks = _split_long_token(draw, word, font, max_width)
            else:
                chunks = [word]
            for chunk in chunks:
                candidate = chunk if not current else f"{current} {chunk}"
                if not current or _measure(draw, candidate, font) <= max_width:
                    current = candidate
                else:
                    lines.append(current)
                    current = chunk
        if current:
            lines.append(current)
    return "\n".join(lines)


def _fit_font_and_text(
    draw: ImageDraw.ImageDraw,
    text: str,
    settings: RenderSettings,
    max_width: int,
    max_height: int,
) -> tuple[object, str, int, tuple[int, int, int, int]]:
    chosen = None
    low, high = settings.font_min_size, settings.font_max_size
    while low <= high:
        size = (low + high) // 2
        font = _load_font(settings.font_path, size)
        wrapped = wrap_text(draw, text, font, max_width)
        spacing = max(4, round(size * settings.line_spacing_ratio))
        bbox = draw.multiline_textbbox(
            (0, 0),
            wrapped,
            font=font,
            spacing=spacing,
            align=settings.quote_position,
            stroke_width=settings.stroke_width,
        )
        if bbox[2] - bbox[0] <= max_width and bbox[3] - bbox[1] <= max_height:
            chosen = (font, wrapped, spacing, bbox)
            low = size + 1
        else:
            high = size - 1
    if chosen is None:
        font = _load_font(settings.font_path, settings.font_min_size)
        wrapped = wrap_text(draw, text, font, max_width)
        spacing = max(4, round(settings.font_min_size * settings.line_spacing_ratio))
        bbox = draw.multiline_textbbox(
            (0, 0),
            wrapped,
            font=font,
            spacing=spacing,
            align=settings.quote_position,
            stroke_width=settings.stroke_width,
        )
        if bbox[3] - bbox[1] > max_height:
            raise QuoteGeneratorError(
                "Quote is too long for the selected text area. "
                "Increase the text area or reduce the minimum font size."
            )
        return font, wrapped, spacing, bbox
    return chosen


def _cover_image(image: Image.Image, settings: RenderSettings) -> Image.Image:
    center_x = {"left": 0.25, "center": 0.5, "right": 0.75}[settings.focal_point]
    return ImageOps.fit(
        ImageOps.exif_transpose(image).convert("RGB"),
        (settings.width, settings.height),
        method=Image.Resampling.LANCZOS,
        centering=(center_x, 0.5),
    )


def _add_vertical_gradient(
    image: Image.Image,
    top_color: tuple[int, int, int, int],
    bottom_color: tuple[int, int, int, int],
) -> Image.Image:
    size = image.size
    mask = Image.linear_gradient("L").resize(size, Image.Resampling.BILINEAR)
    top_layer = Image.new("RGBA", size, top_color)
    bottom_layer = Image.new("RGBA", size, bottom_color)
    gradient = Image.composite(bottom_layer, top_layer, mask)
    return Image.alpha_composite(image.convert("RGBA"), gradient)


def _add_vignette(image: Image.Image, strength: int) -> Image.Image:
    strength = max(0, min(255, strength))
    radial = Image.radial_gradient("L").resize(
        image.size,
        Image.Resampling.BILINEAR,
    )
    alpha = radial.point(lambda value: round(value * strength / 255))
    layer = Image.new("RGBA", image.size, (0, 0, 0, 0))
    layer.putalpha(alpha)
    return Image.alpha_composite(image.convert("RGBA"), layer)


def _add_soft_glow(
    image: Image.Image,
    center: tuple[int, int],
    radius: int,
    color: tuple[int, int, int, int],
) -> Image.Image:
    radius = max(8, radius)
    layer = Image.new("RGBA", image.size, (0, 0, 0, 0))
    glow_draw = ImageDraw.Draw(layer)
    center_x, center_y = center
    glow_draw.ellipse(
        (
            center_x - radius,
            center_y - radius,
            center_x + radius,
            center_y + radius,
        ),
        fill=color,
    )
    layer = layer.filter(ImageFilter.GaussianBlur(max(4, radius // 2)))
    return Image.alpha_composite(image.convert("RGBA"), layer)


def _tiktok_blur_frame(
    source: Image.Image,
    settings: RenderSettings,
) -> Image.Image:
    oriented = ImageOps.exif_transpose(source).convert("RGB")
    minimum = min(settings.width, settings.height)
    background = _cover_image(oriented, settings).filter(
        ImageFilter.GaussianBlur(max(14, round(minimum * 0.025)))
    )
    background = ImageEnhance.Color(background).enhance(0.86)
    background = ImageEnhance.Brightness(background).enhance(0.82)
    background = _add_vertical_gradient(
        background,
        (5, 15, 32, 32),
        (2, 7, 17, 112),
    )

    foreground = ImageOps.contain(
        oriented,
        (
            max(1, round(settings.width * 0.88)),
            max(1, round(settings.height * 0.82)),
        ),
        method=Image.Resampling.LANCZOS,
    )
    foreground = ImageEnhance.Contrast(foreground).enhance(1.04).convert("RGBA")
    foreground = ImageEnhance.Color(foreground).enhance(1.03)
    foreground_width, foreground_height = foreground.size
    left = (settings.width - foreground_width) // 2
    top = (settings.height - foreground_height) // 2
    radius = max(12, round(minimum * 0.025))
    shadow_offset = max(6, round(minimum * 0.012))

    rounded_mask = Image.new("L", foreground.size, 0)
    ImageDraw.Draw(rounded_mask).rounded_rectangle(
        (0, 0, foreground_width - 1, foreground_height - 1),
        radius=radius,
        fill=255,
    )

    shadow = Image.new("RGBA", background.size, (0, 0, 0, 0))
    shadow_draw = ImageDraw.Draw(shadow)
    shadow_draw.rounded_rectangle(
        (
            left,
            top + shadow_offset,
            left + foreground_width,
            top + foreground_height + shadow_offset,
        ),
        radius=radius,
        fill=(0, 0, 0, 170),
    )
    shadow = shadow.filter(
        ImageFilter.GaussianBlur(max(8, round(minimum * 0.018)))
    )

    composed = Image.alpha_composite(background.convert("RGBA"), shadow)
    composed.paste(foreground, (left, top), rounded_mask)

    border = Image.new("RGBA", composed.size, (0, 0, 0, 0))
    ImageDraw.Draw(border).rounded_rectangle(
        (
            left,
            top,
            left + foreground_width - 1,
            top + foreground_height - 1,
        ),
        radius=radius,
        outline=(255, 255, 255, 72),
        width=max(1, round(minimum * 0.002)),
    )
    return Image.alpha_composite(composed, border).convert("RGB")


def _add_luxury_frame(
    canvas: Image.Image,
    settings: RenderSettings,
    width_ratio: float = 0.0015,
    opacity: int = 68,
) -> Image.Image:
    minimum = min(settings.width, settings.height)
    frame = Image.new("RGBA", canvas.size, (0, 0, 0, 0))
    inset = max(10, round(minimum * 0.016))
    ImageDraw.Draw(frame).rounded_rectangle(
        (
            inset,
            inset,
            settings.width - inset - 1,
            settings.height - inset - 1,
        ),
        radius=max(10, round(minimum * 0.018)),
        outline=(218, 178, 96, max(0, min(255, opacity))),
        width=max(1, round(minimum * width_ratio)),
    )
    return Image.alpha_composite(canvas.convert("RGBA"), frame)


def _apply_background_style(
    source: Image.Image,
    settings: RenderSettings,
) -> Image.Image:
    style = settings.background_style
    if style == "tiktok-blur-frame":
        return _tiktok_blur_frame(source, settings)

    canvas = _cover_image(source, settings)
    minimum = min(settings.width, settings.height)

    if style == "tiktok-glass-gradient":
        canvas = ImageEnhance.Color(canvas).enhance(0.94)
        canvas = ImageEnhance.Contrast(canvas).enhance(1.06)
        canvas = _add_soft_glow(
            canvas,
            (round(settings.width * 0.78), round(settings.height * 0.18)),
            round(minimum * 0.42),
            (255, 190, 112, 40),
        )
        canvas = _add_vertical_gradient(
            canvas,
            (8, 24, 48, 22),
            (2, 8, 20, 148),
        )
        canvas = _add_vignette(canvas, 62)
    elif style == "cinematic-vignette":
        canvas = ImageEnhance.Color(canvas).enhance(0.80)
        canvas = ImageEnhance.Contrast(canvas).enhance(1.16)
        canvas = ImageEnhance.Sharpness(canvas).enhance(1.05)
        canvas = _add_vertical_gradient(
            canvas,
            (4, 25, 58, 54),
            (92, 38, 9, 42),
        )
        canvas = _add_vignette(canvas, 118)
    elif style == "luxury-noir":
        canvas = ImageEnhance.Color(canvas).enhance(0.24)
        canvas = ImageEnhance.Contrast(canvas).enhance(1.20)
        canvas = ImageEnhance.Brightness(canvas).enhance(0.92)
        canvas = _add_vertical_gradient(
            canvas,
            (4, 13, 31, 72),
            (1, 5, 15, 118),
        )
        canvas = _add_soft_glow(
            canvas,
            (round(settings.width * 0.76), round(settings.height * 0.22)),
            round(minimum * 0.34),
            (194, 143, 58, 35),
        )
        canvas = _add_vignette(canvas, 148)

        canvas = _add_luxury_frame(canvas, settings)
    elif style == "luxury-noir-color":
        # Keep the source photo's original color treatment and add only the
        # premium gold frame. This variation uses a stronger, thicker border.
        canvas = _add_luxury_frame(
            canvas,
            settings,
            width_ratio=0.0045,
            opacity=150,
        )

    return canvas.convert("RGB")


def _apply_frosted_box(
    canvas: Image.Image,
    box: tuple[int, int, int, int],
    radius: int,
) -> Image.Image:
    size = canvas.size
    left, top, right, bottom = (int(round(value)) for value in box)
    left = max(0, min(size[0] - 1, left))
    top = max(0, min(size[1] - 1, top))
    right = max(left + 1, min(size[0] - 1, right))
    bottom = max(top + 1, min(size[1] - 1, bottom))
    width = max(1, right - left)
    height = max(1, bottom - top)
    region = canvas.crop((left, top, right, bottom)).filter(
        ImageFilter.GaussianBlur(max(5, round(min(width, height) * 0.035)))
    )
    mask = Image.new("L", (width, height), 0)
    ImageDraw.Draw(mask).rounded_rectangle(
        (0, 0, width - 1, height - 1),
        radius=min(radius, width // 2, height // 2),
        fill=255,
    )
    result = canvas.copy()
    result.paste(region, (left, top), mask)
    return result


def _create_quote_frame_layer(
    size: tuple[int, int],
    box: tuple[int, int, int, int],
    settings: RenderSettings,
) -> Image.Image:
    left, top, right, bottom = (int(round(value)) for value in box)
    left = max(0, min(size[0] - 1, left))
    top = max(0, min(size[1] - 1, top))
    right = max(left + 1, min(size[0] - 1, right))
    bottom = max(top + 1, min(size[1] - 1, bottom))
    minimum = min(size)
    line_width = max(3, round(minimum * 0.006))
    radius = max(line_width * 2, round(minimum * 0.028))
    quote_size = max(38, round(minimum * 0.075))
    frame_alpha = min(255, round(settings.box_opacity * 1.7))
    frame_color = _hex_to_rgba(settings.box_color, frame_alpha)

    layer = Image.new("RGBA", size, (0, 0, 0, 0))
    frame_draw = ImageDraw.Draw(layer)
    frame_draw.rounded_rectangle(
        (left, top, right, bottom),
        radius=radius,
        outline=frame_color,
        width=line_width,
    )

    quote_font = _load_font("", quote_size)
    opening = "\u201c"
    closing = "\u201d"
    opening_bbox = frame_draw.textbbox((0, 0), opening, font=quote_font)
    closing_bbox = frame_draw.textbbox((0, 0), closing, font=quote_font)
    opening_width = opening_bbox[2] - opening_bbox[0]
    opening_height = opening_bbox[3] - opening_bbox[1]
    closing_width = closing_bbox[2] - closing_bbox[0]
    closing_height = closing_bbox[3] - closing_bbox[1]
    gap_padding = max(line_width * 2, round(minimum * 0.009))

    opening_left = left + line_width
    opening_top = top + line_width
    closing_left = right - line_width - closing_width
    closing_top = bottom - line_width - closing_height

    frame_draw.rectangle(
        (
            max(0, left - line_width),
            max(0, top - line_width),
            min(size[0], opening_left + opening_width + gap_padding),
            min(size[1], opening_top + opening_height + gap_padding),
        ),
        fill=(0, 0, 0, 0),
    )
    frame_draw.rectangle(
        (
            max(0, closing_left - gap_padding),
            max(0, closing_top - gap_padding),
            min(size[0], right + line_width),
            min(size[1], bottom + line_width),
        ),
        fill=(0, 0, 0, 0),
    )

    frame_draw.text(
        (
            opening_left - opening_bbox[0],
            opening_top - opening_bbox[1],
        ),
        opening,
        font=quote_font,
        fill=frame_color,
    )
    frame_draw.text(
        (
            closing_left - closing_bbox[0],
            closing_top - closing_bbox[1],
        ),
        closing,
        font=quote_font,
        fill=frame_color,
    )
    return layer


def _create_corner_quote_frame_layer(
    size: tuple[int, int],
    box: tuple[int, int, int, int],
    settings: RenderSettings,
) -> Image.Image:
    left, top, right, bottom = (int(round(value)) for value in box)
    left = max(0, min(size[0] - 1, left))
    top = max(0, min(size[1] - 1, top))
    right = max(left + 1, min(size[0] - 1, right))
    bottom = max(top + 1, min(size[1] - 1, bottom))
    minimum = min(size)
    frame_width = right - left
    frame_height = bottom - top
    line_width = max(3, round(minimum * 0.006))
    horizontal_length = min(
        round(frame_width * 0.28),
        max(line_width * 5, round(minimum * 0.16)),
    )
    vertical_length = min(
        round(frame_height * 0.28),
        max(line_width * 5, round(minimum * 0.16)),
    )
    quote_size = max(42, round(minimum * 0.085))
    frame_alpha = min(255, round(settings.box_opacity * 1.7))
    frame_color = _hex_to_rgba(settings.box_color, frame_alpha)

    layer = Image.new("RGBA", size, (0, 0, 0, 0))
    frame_draw = ImageDraw.Draw(layer)
    corners = (
        (
            (left + horizontal_length, top),
            (left, top),
            (left, top + vertical_length),
        ),
        (
            (right - horizontal_length, top),
            (right, top),
            (right, top + vertical_length),
        ),
        (
            (left, bottom - vertical_length),
            (left, bottom),
            (left + horizontal_length, bottom),
        ),
        (
            (right - horizontal_length, bottom),
            (right, bottom),
            (right, bottom - vertical_length),
        ),
    )
    for points in corners:
        frame_draw.line(
            points,
            fill=frame_color,
            width=line_width,
            joint="curve",
        )

    quote_font = _load_font("", quote_size)
    for mark, vertical in (("\u201c", "top"), ("\u201d", "bottom")):
        mark_bbox = frame_draw.textbbox((0, 0), mark, font=quote_font)
        mark_width = mark_bbox[2] - mark_bbox[0]
        mark_height = mark_bbox[3] - mark_bbox[1]
        mark_x = (left + right - mark_width) // 2 - mark_bbox[0]
        if vertical == "top":
            mark_y = top + line_width - mark_bbox[1]
        else:
            mark_y = bottom - line_width - mark_height - mark_bbox[1]
        frame_draw.text(
            (mark_x, mark_y),
            mark,
            font=quote_font,
            fill=frame_color,
        )
    return layer


def _create_double_line_frame_layer(
    size: tuple[int, int],
    box: tuple[int, int, int, int],
    settings: RenderSettings,
) -> Image.Image:
    left, top, right, bottom = (int(round(value)) for value in box)
    left = max(0, min(size[0] - 1, left))
    top = max(0, min(size[1] - 1, top))
    right = max(left + 1, min(size[0] - 1, right))
    bottom = max(top + 1, min(size[1] - 1, bottom))
    minimum = min(size)
    line_width = max(2, round(minimum * 0.0035))
    offset = max(line_width * 3, round(minimum * 0.018))
    offset = min(offset, (right - left) // 4, (bottom - top) // 4)
    frame_alpha = min(255, round(settings.box_opacity * 1.7))
    frame_color = _hex_to_rgba(settings.box_color, frame_alpha)

    layer = Image.new("RGBA", size, (0, 0, 0, 0))
    frame_draw = ImageDraw.Draw(layer)
    frame_draw.rectangle(
        (
            left + offset,
            top,
            right,
            bottom - offset,
        ),
        outline=frame_color,
        width=line_width,
    )
    frame_draw.rectangle(
        (
            left,
            top + offset,
            right - offset,
            bottom,
        ),
        outline=frame_color,
        width=line_width,
    )
    return layer


def _unique_output_path(
    output_folder: Path,
    index: int,
    quote: str,
    output_format: str,
) -> Path:
    slug = re.sub(r"[^A-Za-z0-9]+", "_", quote).strip("_")[:42] or "quote"
    extension = ".png" if output_format.upper() == "PNG" else ".jpg"
    base = output_folder / f"quote_{index:03d}_{slug}{extension}"
    if not base.exists():
        return base
    counter = 2
    while True:
        candidate = output_folder / f"quote_{index:03d}_{slug}_{counter}{extension}"
        if not candidate.exists():
            return candidate
        counter += 1


def render_quote_image(
    image_path_value: str | Path,
    quote: str,
    settings: RenderSettings,
) -> Image.Image:
    settings.validate()
    image_path = Path(image_path_value)
    if not image_path.is_file():
        raise QuoteGeneratorError(f"Source image does not exist: {image_path}")
    cleaned_quote = _clean_quote(quote)
    if not cleaned_quote:
        raise QuoteGeneratorError("Quote cannot be empty.")

    try:
        with Image.open(image_path) as source:
            canvas = _apply_background_style(source, settings)
    except (OSError, ValueError) as exc:
        raise QuoteGeneratorError(f"Unable to open image: {image_path.name}") from exc

    if settings.blur_radius > 0:
        canvas = canvas.filter(ImageFilter.GaussianBlur(settings.blur_radius))
    if settings.darken_percent:
        brightness = max(0.05, 1.0 - settings.darken_percent / 100.0)
        canvas = ImageEnhance.Brightness(canvas).enhance(brightness)

    canvas = canvas.convert("RGBA")
    overlay = Image.new("RGBA", canvas.size, (0, 0, 0, 0))
    draw = ImageDraw.Draw(overlay)

    margin = round(min(settings.width, settings.height) * settings.safe_margin_percent / 100)
    area_width = round(settings.width * settings.text_area_percent / 100)
    if settings.quote_position == "center":
        area_width = min(area_width, settings.width - 2 * margin)
        area_left = (settings.width - area_width) // 2
    elif settings.quote_position == "right":
        area_left = settings.width - margin - area_width
    else:
        area_left = margin
    area_right = area_left + area_width
    max_text_height = round(settings.height * 0.62)

    decorative_frame = (
        settings.box_enabled
        and settings.box_opacity > 0
        and settings.box_style in {"quote-frame", "corner-frame"}
    )
    visible_quote = (
        cleaned_quote
        if decorative_frame
        else (
            f"\u201c{cleaned_quote}\u201d"
            if settings.add_quote_marks
            else cleaned_quote
        )
    )
    font, wrapped, spacing, bbox = _fit_font_and_text(
        draw,
        visible_quote,
        settings,
        max_width=area_width,
        max_height=max_text_height,
    )
    text_width = bbox[2] - bbox[0]
    text_height = bbox[3] - bbox[1]

    if settings.quote_position == "center":
        text_x = (settings.width - text_width) // 2
    elif settings.quote_position == "right":
        text_x = area_right - text_width
    else:
        text_x = area_left

    desired_center_y = round(settings.height * settings.vertical_position / 100)
    text_y = desired_center_y - text_height // 2 - bbox[1]
    min_y = margin - bbox[1]
    max_y = settings.height - margin - text_height - bbox[1]
    text_y = max(min_y, min(text_y, max_y))

    padding = round(min(settings.width, settings.height) * settings.box_padding_percent / 100)
    radius = round(min(settings.width, settings.height) * settings.box_radius_percent / 100)
    visual_left = text_x + bbox[0]
    visual_top = text_y + bbox[1]
    visual_right = visual_left + text_width
    visual_bottom = visual_top + text_height

    if settings.box_enabled and settings.box_opacity > 0:
        if settings.box_style == "corner-frame":
            active_padding = max(
                padding,
                round(min(settings.width, settings.height) * 0.085),
            )
        elif settings.box_style == "quote-frame":
            active_padding = max(
                padding,
                round(min(settings.width, settings.height) * 0.050),
            )
        elif settings.box_style == "double-frame":
            active_padding = max(
                padding,
                round(min(settings.width, settings.height) * 0.055),
            )
        else:
            active_padding = padding
        box = (
            max(0, visual_left - active_padding),
            max(0, visual_top - active_padding),
            min(settings.width - 1, visual_right + active_padding),
            min(settings.height - 1, visual_bottom + active_padding),
        )
        if settings.box_style == "quote-frame":
            frame_layer = _create_quote_frame_layer(
                overlay.size,
                box,
                settings,
            )
            overlay = Image.alpha_composite(overlay, frame_layer)
            draw = ImageDraw.Draw(overlay)
        elif settings.box_style == "corner-frame":
            frame_layer = _create_corner_quote_frame_layer(
                overlay.size,
                box,
                settings,
            )
            overlay = Image.alpha_composite(overlay, frame_layer)
            draw = ImageDraw.Draw(overlay)
        elif settings.box_style == "double-frame":
            frame_layer = _create_double_line_frame_layer(
                overlay.size,
                box,
                settings,
            )
            overlay = Image.alpha_composite(overlay, frame_layer)
            draw = ImageDraw.Draw(overlay)
        else:
            if settings.background_style == "tiktok-glass-gradient":
                canvas = _apply_frosted_box(canvas, box, radius)
            draw.rounded_rectangle(
                box,
                radius=radius,
                fill=_hex_to_rgba(settings.box_color, settings.box_opacity),
                outline=(
                    (255, 255, 255, 54)
                    if settings.background_style == "tiktok-glass-gradient"
                    else None
                ),
                width=max(
                    1,
                    round(min(settings.width, settings.height) * 0.0015),
                ),
            )

    if settings.shadow_enabled:
        shadow_fill = _hex_to_rgba(settings.shadow_color, settings.shadow_opacity)
        draw.multiline_text(
            (text_x + settings.shadow_offset, text_y + settings.shadow_offset),
            wrapped,
            font=font,
            fill=shadow_fill,
            spacing=spacing,
            align=settings.quote_position,
            stroke_width=settings.stroke_width + 1,
            stroke_fill=shadow_fill,
        )

    text_fill = _hex_to_rgba(settings.text_color, 255)
    stroke_fill = _hex_to_rgba(settings.shadow_color, 220)
    draw.multiline_text(
        (text_x, text_y),
        wrapped,
        font=font,
        fill=text_fill,
        spacing=spacing,
        align=settings.quote_position,
        stroke_width=settings.stroke_width,
        stroke_fill=stroke_fill,
    )

    watermark = settings.watermark.strip()
    if watermark:
        wm_size = max(18, round(settings.width * 0.018))
        wm_font = _load_font("", wm_size, bold_fallback=False)
        wm_bbox = draw.textbbox((0, 0), watermark, font=wm_font)
        wm_width = wm_bbox[2] - wm_bbox[0]
        wm_height = wm_bbox[3] - wm_bbox[1]
        wm_margin = max(20, round(min(settings.width, settings.height) * 0.035))
        left_x = wm_margin - wm_bbox[0]
        center_x = (settings.width - wm_width) // 2 - wm_bbox[0]
        right_x = settings.width - wm_margin - wm_width - wm_bbox[0]
        top_y = wm_margin - wm_bbox[1]
        bottom_y = settings.height - wm_margin - wm_height - wm_bbox[1]
        positions = {
            "top-left": (left_x, top_y),
            "top-center": (center_x, top_y),
            "top-right": (right_x, top_y),
            "bottom-left": (left_x, bottom_y),
            "bottom-center": (center_x, bottom_y),
            "bottom-right": (right_x, bottom_y),
        }
        wm_x, wm_y = positions.get(
            settings.watermark_position, positions["bottom-right"]
        )
        wm_shadow_offset = max(1, round(settings.width * 0.001))
        draw.text(
            (wm_x + wm_shadow_offset, wm_y + wm_shadow_offset),
            watermark,
            font=wm_font,
            fill=(0, 0, 0, min(150, settings.watermark_opacity)),
        )
        draw.text(
            (wm_x, wm_y),
            watermark,
            font=wm_font,
            fill=_hex_to_rgba(settings.watermark_color, settings.watermark_opacity),
        )

    return Image.alpha_composite(canvas, overlay).convert("RGB")


def generate_batch(
    items: Sequence[GenerationItem],
    output_folder_value: str | Path,
    settings: RenderSettings,
    *,
    progress_callback: Callable[[int, int, GenerationResult], None] | None = None,
    should_cancel: Callable[[], bool] | None = None,
    create_timestamp_folder: bool = True,
) -> tuple[Path, list[GenerationResult]]:
    settings.validate()
    output_root = Path(output_folder_value)
    output_root.mkdir(parents=True, exist_ok=True)
    if create_timestamp_folder:
        stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_folder = output_root / f"Quote_Images_{stamp}"
        suffix = 2
        while output_folder.exists():
            output_folder = output_root / f"Quote_Images_{stamp}_{suffix}"
            suffix += 1
        output_folder.mkdir(parents=True)
    else:
        output_folder = output_root

    results: list[GenerationResult] = []
    total = len(items)
    for position, item in enumerate(items, start=1):
        if should_cancel and should_cancel():
            break
        rendered = render_quote_image(item.image_path, item.quote, settings)
        output_path = _unique_output_path(
            output_folder,
            item.index,
            item.quote,
            settings.output_format,
        )
        if settings.output_format.upper() == "PNG":
            rendered.save(output_path, format="PNG", optimize=True)
        else:
            rendered.save(
                output_path,
                format="JPEG",
                quality=settings.jpeg_quality,
                optimize=True,
                subsampling=0,
            )
        result = GenerationResult(
            output_path=output_path,
            image_path=item.image_path,
            quote=item.quote,
            index=item.index,
        )
        results.append(result)
        if progress_callback:
            progress_callback(position, total, result)

    log_path = output_folder / "generation_log.csv"
    with log_path.open("w", newline="", encoding="utf-8-sig") as handle:
        writer = csv.writer(handle)
        writer.writerow(["No.", "Output file", "Source image", "Quote"])
        for result in results:
            writer.writerow(
                [
                    result.index,
                    result.output_path.name,
                    str(result.image_path),
                    result.quote,
                ]
            )
    return output_folder, results
